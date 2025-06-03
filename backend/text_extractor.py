import PyPDF2
import pdfplumber
from docx import Document
import pytesseract
from PIL import Image
import io
import magic
import re
from typing import Union, Optional, Dict, Any, List, Tuple
import logging
from collections import Counter
import email
import phonenumbers
from urllib.parse import urlparse

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class TextExtractor:
    def __init__(self):
        self.supported_formats = {
            'pdf': ['.pdf'],
            'docx': ['.docx', '.doc'],
            'image': ['.png', '.jpg', '.jpeg', '.tiff', '.bmp'],
            'text': ['.txt'],
            'html': ['.html', '.htm'],
            'rtf': ['.rtf']
        }
        self.patterns = {
            'email': re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
            'phone': re.compile(r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'),
            'linkedin': re.compile(r'(?:linkedin\.com/in/|linkedin\.com/pub/)([a-zA-Z0-9\-_%]+)', re.IGNORECASE),
            'github': re.compile(r'(?:github\.com/)([a-zA-Z0-9\-_]+)', re.IGNORECASE),
            'website': re.compile(r'https?://(?:[-\w.])+(?:\:[0-9]+)?(?:/(?:[\w/_.])*(?:\?(?:[\w&=%.])*)?(?:#(?:[\w.])*)?)?'),
            'degree': re.compile(r'\b(?:Bachelor|Master|PhD|Ph\.D|MBA|M\.S\.|B\.S\.|B\.A\.|M\.A\.|Doctor|Associates?|A\.S\.|A\.A\.)\b', re.IGNORECASE),
            'gpa': re.compile(r'GPA:?\s*([0-3]?\.[0-9]{1,2}|[0-4]\.[0-9]{1,2})', re.IGNORECASE),
            'year': re.compile(r'\b(19|20)\d{2}\b'),
            'date_range': re.compile(r'(\d{1,2}[\/\-]\d{4}|\w+\s+\d{4}|\d{4})\s*[-–—to]\s*(\d{1,2}[\/\-]\d{4}|\w+\s+\d{4}|\d{4}|present|current)', re.IGNORECASE),
            'salary': re.compile(r'\$\s*([0-9,]+)\s*(?:k|thousand)?(?:\s*[-–—]\s*\$?\s*([0-9,]+)\s*(?:k|thousand)?)?', re.IGNORECASE),
            'achievements': re.compile(r'(?:achieved|increased|decreased|improved|reduced|implemented|developed|led|managed|created|built|designed|optimized)\s+[^.!?]*(?:\d+%?|million|thousand|k|\$[0-9,]+)', re.IGNORECASE),
            'location': re.compile(r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*([A-Z]{2}|[A-Z][a-z]+)\b'),
            'institution': re.compile(r'\b(?:University|College|Institute|School)\s+of\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*|\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:University|College|Institute|School)\b', re.IGNORECASE),
            'company': re.compile(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Inc|LLC|Corp|Corporation|Company|Technologies|Systems|Solutions|Group|Consulting)\b', re.IGNORECASE)
        }
        
        self.section_headers = {
            'personal_info': {
                'patterns': ['contact', 'personal information', 'about me', 'personal details'],
                'confidence': 0.9
            },
            'summary': {
                'patterns': ['summary', 'profile', 'objective', 'professional summary',
                           'career objective', 'personal statement', 'about', 'overview'],
                'confidence': 0.95
            },
            'experience': {
                'patterns': ['experience', 'work experience', 'employment', 'professional experience',
                           'career history', 'work history', 'positions', 'employment history',
                           'professional background', 'work', 'career'],
                'confidence': 0.98
            },
            'education': {
                'patterns': ['education', 'academic background', 'educational background',
                           'qualifications', 'academic qualifications', 'degrees', 'academic'],
                'confidence': 0.95
            },
            'skills': {
                'patterns': ['skills', 'technical skills', 'core competencies', 'expertise',
                           'proficiencies', 'technologies', 'tools', 'programming languages',
                           'technical competencies', 'software', 'languages'],
                'confidence': 0.92
            },
            'projects': {
                'patterns': ['projects', 'key projects', 'notable projects', 'personal projects',
                           'academic projects', 'portfolio', 'selected projects'],
                'confidence': 0.88
            },
            'certifications': {
                'patterns': ['certifications', 'certificates', 'professional certifications',
                           'licenses', 'credentials', 'training'],
                'confidence': 0.90
            },
            'awards': {
                'patterns': ['awards', 'honors', 'achievements', 'recognition', 'accomplishments'],
                'confidence': 0.85
            },
            'requirements': {
                'patterns': ['requirements', 'qualifications', 'required', 'must have',
                           'minimum requirements', 'prerequisites'],
                'confidence': 0.95
            },
            'responsibilities': {
                'patterns': ['responsibilities', 'duties', 'job description', 'role',
                           'key responsibilities', 'main duties'],
                'confidence': 0.93
            },
            'benefits': {
                'patterns': ['benefits', 'compensation', 'salary', 'package', 'perks'],
                'confidence': 0.88
            }
        }
    
    def extract_text(self, file_path: str, file_content: bytes = None) -> Dict[str, Any]:
        try:
            file_type = self._detect_file_type(file_path, file_content)
            
            result = {
                'raw_text': '',
                'cleaned_text': '',
                'structured_data': {},
                'sections': {},
                'contact_info': {},
                'file_type': file_type,
                'extraction_method': '',
                'success': False,
                'metadata': {},
                'quality_score': 0.0,
                'keywords': [],
                'achievements': [],
                'document_insights': {}
            }
            
            if file_type == 'pdf':
                raw_result = self._extract_from_pdf(file_path, file_content)
            elif file_type == 'docx':
                raw_result = self._extract_from_docx(file_path, file_content)
            elif file_type == 'image':
                raw_result = self._extract_from_image(file_path, file_content)
            elif file_type == 'text':
                raw_result = self._extract_from_text(file_path, file_content)
            else:
                raw_result = {
                    'text': "Unsupported file format",
                    'success': False,
                    'extraction_method': 'unsupported'
                }
            
            if not raw_result.get('success', False):
                return raw_result
            
            raw_text = raw_result.get('text', '')
            result['raw_text'] = raw_text
            result['cleaned_text'] = self._clean_text(raw_text)
            result['extraction_method'] = raw_result.get('extraction_method', '')
            result['success'] = True
            result['metadata'] = raw_result.get('metadata', {})
            
            result['contact_info'] = self._extract_contact_info(raw_text)
            result['sections'] = self._extract_sections(raw_text)
            result['structured_data'] = self._extract_structured_data(raw_text)
            result['keywords'] = self._extract_keywords(raw_text)
            result['achievements'] = self._extract_achievements(raw_text)
            result['document_insights'] = self._generate_document_insights(result)
            result['quality_score'] = self._calculate_quality_score(result)
            
            result['text'] = result['cleaned_text']
            
            return result
            
        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            return {
                'raw_text': '',
                'cleaned_text': '',
                'text': '',
                'structured_data': {},
                'sections': {},
                'contact_info': {},
                'file_type': 'unknown',
                'extraction_method': 'failed',
                'success': False,
                'error': str(e),
                'metadata': {},
                'quality_score': 0.0,
                'keywords': [],
                'achievements': [],
                'document_insights': {}
            }
    
    def _detect_file_type(self, file_path: str, file_content: bytes = None) -> str:
        try:
            extension = file_path.lower().split('.')[-1] if '.' in file_path else ''
            
            for file_type, extensions in self.supported_formats.items():
                if f'.{extension}' in extensions:
                    return file_type
            
            if file_content:
                mime_type = magic.from_buffer(file_content, mime=True)
                if 'pdf' in mime_type:
                    return 'pdf'
                elif 'word' in mime_type or 'officedocument' in mime_type:
                    return 'docx'
                elif 'image' in mime_type:
                    return 'image'
                elif 'text' in mime_type:
                    return 'text'
            
            return 'unknown'
            
        except Exception as e:
            logger.warning(f"File type detection failed: {str(e)}")
            return 'unknown'
    
    def _extract_from_pdf(self, file_path: str, file_content: bytes = None) -> Dict[str, Any]:
        methods_tried = []
        
        try:
            if file_content:
                pdf_file = io.BytesIO(file_content)
            else:
                pdf_file = file_path
                
            with pdfplumber.open(pdf_file) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                if text_parts:
                    methods_tried.append('pdfplumber')
                    return {
                        'text': '\n'.join(text_parts),
                        'file_type': 'pdf',
                        'extraction_method': 'pdfplumber',
                        'success': True,
                        'metadata': {
                            'pages': len(pdf.pages),
                            'methods_tried': methods_tried
                        }
                    }
        except Exception as e:
            methods_tried.append(f'pdfplumber_failed: {str(e)}')
            logger.warning(f"pdfplumber extraction failed: {str(e)}")
        
        try:
            if file_content:
                pdf_file = io.BytesIO(file_content)
            else:
                pdf_file = open(file_path, 'rb')
                
            reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            if not file_content:
                pdf_file.close()
                
            if text_parts:
                methods_tried.append('PyPDF2')
                return {
                    'text': '\n'.join(text_parts),
                    'file_type': 'pdf',
                    'extraction_method': 'PyPDF2',
                    'success': True,
                    'metadata': {
                        'pages': len(reader.pages),
                        'methods_tried': methods_tried
                    }
                }
                
        except Exception as e:
            methods_tried.append(f'PyPDF2_failed: {str(e)}')
            logger.warning(f"PyPDF2 extraction failed: {str(e)}")
        
        try:
            text = self._ocr_pdf(file_path, file_content)
            if text.strip():
                methods_tried.append('OCR')
                return {
                    'text': text,
                    'file_type': 'pdf',
                    'extraction_method': 'OCR',
                    'success': True,
                    'metadata': {
                        'methods_tried': methods_tried,
                        'note': 'Extracted using OCR - may contain errors'
                    }
                }
        except Exception as e:
            methods_tried.append(f'OCR_failed: {str(e)}')
            logger.warning(f"OCR extraction failed: {str(e)}")
        
        return {
            'text': '',
            'file_type': 'pdf',
            'extraction_method': 'failed',
            'success': False,
            'metadata': {'methods_tried': methods_tried}
        }
    
    def _extract_from_docx(self, file_path: str, file_content: bytes = None) -> Dict[str, Any]:
        try:
            if file_content:
                doc_file = io.BytesIO(file_content)
            else:
                doc_file = file_path
                
            doc = Document(doc_file)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return {
                'text': '\n'.join(text_parts),
                'file_type': 'docx',
                'extraction_method': 'python-docx',
                'success': True,
                'metadata': {
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables)
                }
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return {
                'text': '',
                'file_type': 'docx',
                'extraction_method': 'failed',
                'success': False,
                'error': str(e),
                'metadata': {}
            }
    
    def _extract_from_image(self, file_path: str, file_content: bytes = None) -> Dict[str, Any]:
        try:
            if file_content:
                image = Image.open(io.BytesIO(file_content))
            else:
                image = Image.open(file_path)
            
            custom_config = r'--oem 3 --psm 6'
            text = pytesseract.image_to_string(image, config=custom_config)
            
            return {
                'text': text,
                'file_type': 'image',
                'extraction_method': 'pytesseract',
                'success': True,
                'metadata': {
                    'image_size': image.size,
                    'image_mode': image.mode
                }
            }
            
        except Exception as e:
            logger.error(f"Image OCR extraction failed: {str(e)}")
            return {
                'text': '',
                'file_type': 'image',
                'extraction_method': 'failed',
                'success': False,
                'error': str(e),
                'metadata': {}
            }
    
    def _extract_from_text(self, file_path: str, file_content: bytes = None) -> Dict[str, Any]:
        try:
            if file_content:
                text = file_content.decode('utf-8')
            else:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            
            return {
                'text': text,
                'file_type': 'text',
                'extraction_method': 'direct',
                'success': True,
                'metadata': {
                    'character_count': len(text),
                    'line_count': len(text.split('\n'))
                }
            }
            
        except UnicodeDecodeError:
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    if file_content:
                        text = file_content.decode(encoding)
                    else:
                        with open(file_path, 'r', encoding=encoding) as file:
                            text = file.read()
                    
                    return {
                        'text': text,
                        'file_type': 'text',
                        'extraction_method': f'direct_{encoding}',
                        'success': True,
                        'metadata': {
                            'encoding_used': encoding,
                            'character_count': len(text)
                        }
                    }
                except:
                    continue
                    
        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            return {
                'text': '',
                'file_type': 'text',
                'extraction_method': 'failed',
                'success': False,
                'error': str(e),
                'metadata': {}
            }
    
    def _ocr_pdf(self, file_path: str, file_content: bytes = None) -> str:
        try:
            import fitz  # PyMuPDF for PDF to image conversion
            
            if file_content:
                doc = fitz.open(stream=file_content, filetype="pdf")
            else:
                doc = fitz.open(file_path)
            
            text_parts = []
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                
                page_text = pytesseract.image_to_string(img)
                if page_text.strip():
                    text_parts.append(page_text)
            
            doc.close()
            return '\n'.join(text_parts)
            
        except ImportError:
            logger.warning("PyMuPDF not available for PDF OCR")
            return ""
        except Exception as e:
            logger.error(f"PDF OCR failed: {str(e)}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        if not text:
            return ""
        
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n[ \t]+', '\n', text)
        
        text = re.sub(r'[^\w\s\n\.\-@#\(\)\[\]{}:;,!?\'"/$%&*+=<>|\\`~]', ' ', text)
        
        text = re.sub(r'\bl\b', 'I', text)
        text = re.sub(r'\b0\b', 'O', text)
        text = re.sub(r'(\w)\s+([.,!?;:])', r'\1\2', text)
        
        text = re.sub(r'\s+([,.!?;:])', r'\1', text)
        text = re.sub(r'([,.!?;:])\s*([a-zA-Z])', r'\1 \2', text)
        
        text = re.sub(r'^[\s•·▪▫◦‣⁃]\s*', '• ', text, flags=re.MULTILINE)
        
        text = re.sub(r'(\d{1,2})/(\d{4})', r'\1/\2', text)
        text = re.sub(r'(\w{3,9})\s+(\d{4})', r'\1 \2', text)
        
        return text.strip()
    
    def _extract_contact_info(self, text: str) -> Dict[str, Any]:
        contact_info = {
            'emails': [],
            'phones': [],
            'linkedin': None,
            'github': None,
            'websites': [],
            'location': None,
            'name': None
        }
        
        emails = self.patterns['email'].findall(text)
        contact_info['emails'] = list(set(emails))
        
        phone_matches = self.patterns['phone'].findall(text)
        phones = []
        for match in phone_matches:
            if isinstance(match, tuple):
                area_code = match[1] if match[1] else ''
                prefix = match[2] if match[2] else ''
                number = match[3] if match[3] else ''
                if area_code and prefix and number:
                    formatted_phone = f"({area_code}) {prefix}-{number}"
                    phones.append(formatted_phone)
            else:
                clean_phone = re.sub(r'[^\d]', '', str(match))
                if len(clean_phone) >= 10:
                    phones.append(str(match))
        contact_info['phones'] = list(set(phones))
        
        linkedin_match = self.patterns['linkedin'].search(text)
        if linkedin_match:
            contact_info['linkedin'] = f"https://{linkedin_match.group(0)}"
        
        github_match = self.patterns['github'].search(text)
        if github_match:
            contact_info['github'] = f"https://{github_match.group(0)}"
        
        websites = self.patterns['website'].findall(text)
        contact_info['websites'] = list(set(websites))
        
        location_match = self.patterns['location'].search(text)
        if location_match:
            contact_info['location'] = location_match.group(0)
        
        lines = text.split('\n')[:5]
        for line in lines:
            line = line.strip()
            if len(line) < 50 and len(line.split()) <= 4:
                words = line.split()
                if 2 <= len(words) <= 4 and all(word[0].isupper() for word in words if word.isalpha()):
                    contact_info['name'] = line
                    break
        
        return contact_info
    
    def _extract_sections(self, text: str) -> Dict[str, str]:
        sections = {}
        lines = text.split('\n')
        current_section = 'general'
        current_content = []
        section_confidence = {}
        
        for i, line in enumerate(lines):
            line_clean = line.strip()
            line_lower = line_clean.lower()
            section_found = None
            best_confidence = 0
            
            if len(line_clean) < 3 or len(line_clean) > 80:
                current_content.append(line)
                continue
            
            for section_type, section_data in self.section_headers.items():
                patterns = section_data['patterns']
                base_confidence = section_data['confidence']
                
                for pattern in patterns:
                    if pattern in line_lower:
                        confidence = base_confidence
                        
                        if len(line_clean) < 30:
                            confidence += 0.05
                        
                        if line_clean.isupper() or line_clean.istitle():
                            confidence += 0.03
                        
                        if ':' in line_clean:
                            confidence += 0.02
                        
                        next_lines = lines[i+1:i+4] if i+1 < len(lines) else []
                        if any(len(next_line.strip()) > 20 for next_line in next_lines):
                            confidence += 0.02
                        
                        if confidence > best_confidence:
                            section_found = section_type
                            best_confidence = confidence
            
            if section_found and best_confidence > 0.85:
                if current_content:
                    content = '\n'.join(current_content).strip()
                    if content:
                        sections[current_section] = content
                        section_confidence[current_section] = section_confidence.get(current_section, 0.5)
                
                current_section = section_found
                section_confidence[current_section] = best_confidence
                current_content = []
            else:
                current_content.append(line)
        
        if current_content:
            content = '\n'.join(current_content).strip()
            if content:
                sections[current_section] = content
        
        for section_name in sections:
            if section_name not in section_confidence:
                section_confidence[section_name] = 0.5
        
        if section_confidence:
            sections['_metadata'] = {
                'section_confidence': section_confidence,
                'total_sections': len(sections) - 1
            }
        
        return sections
    
    def _extract_structured_data(self, text: str) -> Dict[str, Any]:
        structured_data = {
            'education': {
                'degrees': [],
                'institutions': [],
                'gpa': None,
                'graduation_years': []
            },
            'experience': {
                'companies': [],
                'positions': [],
                'years_experience': [],
                'date_ranges': []
            },
            'technical_info': {
                'programming_languages': [],
                'frameworks': [],
                'tools': [],
                'certifications': []
            },
            'metrics': {
                'salary_mentioned': False,
                'salary_range': None,
                'years_total_experience': None
            }
        }
        
        degree_matches = self.patterns['degree'].findall(text)
        structured_data['education']['degrees'] = list(set(degree_matches))
        
        gpa_match = self.patterns['gpa'].search(text)
        if gpa_match:
            try:
                structured_data['education']['gpa'] = float(gpa_match.group(1))
            except (ValueError, IndexError):
                pass
        
        years = self.patterns['year'].findall(text)
        structured_data['education']['graduation_years'] = sorted(set(int(year) for year in years))
        
        date_ranges = self.patterns['date_range'].findall(text)
        structured_data['experience']['date_ranges'] = date_ranges
        
        salary_matches = self.patterns['salary'].findall(text)
        if salary_matches:
            structured_data['metrics']['salary_mentioned'] = True
            structured_data['metrics']['salary_range'] = salary_matches[0] if salary_matches[0] else None
        
        tech_patterns = {
            'programming_languages': [
                'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php',
                'swift', 'kotlin', 'go', 'rust', 'scala', 'r', 'matlab', 'sql'
            ],
            'frameworks': [
                'react', 'angular', 'vue', 'django', 'flask', 'spring', 'express',
                'laravel', 'rails', 'asp.net', 'tensorflow', 'pytorch'
            ],
            'tools': [
                'git', 'docker', 'kubernetes', 'jenkins', 'aws', 'azure', 'gcp',
                'linux', 'windows', 'macos', 'mysql', 'postgresql', 'mongodb'
            ]
        }
        
        text_lower = text.lower()
        for category, terms in tech_patterns.items():
            found_terms = []
            for term in terms:
                if term in text_lower:
                    found_terms.append(term)
            structured_data['technical_info'][category] = found_terms
        
        return structured_data
    
    def _extract_keywords(self, text: str) -> List[Dict[str, Any]]:
        keywords = {
            'technical_skills': [],
            'soft_skills': [],
            'industries': [],
            'job_titles': [],
            'action_words': []
        }
        
        text_lower = text.lower()
        
        technical_terms = {
            'programming': ['python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin', 'go', 'rust', 'scala', 'r', 'matlab', 'sql', 'html', 'css'],
            'frameworks': ['react', 'angular', 'vue', 'django', 'flask', 'spring', 'express', 'laravel', 'rails', 'asp.net', 'tensorflow', 'pytorch', 'keras', 'node.js'],
            'databases': ['mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'cassandra', 'oracle', 'sqlite'],
            'cloud': ['aws', 'azure', 'gcp', 'google cloud', 'amazon web services', 'microsoft azure'],
            'tools': ['git', 'docker', 'kubernetes', 'jenkins', 'linux', 'windows', 'macos', 'jira', 'confluence']
        }
        
        for category, terms in technical_terms.items():
            for term in terms:
                if term in text_lower:
                    confidence = 0.7
                    if f' {term} ' in text_lower:
                        confidence += 0.2
                    if term in text_lower.split():
                        confidence += 0.1
                    
                    keywords['technical_skills'].append({
                        'term': term,
                        'category': category,
                        'confidence': min(confidence, 1.0)
                    })
        
        soft_skills_terms = [
            'leadership', 'communication', 'teamwork', 'problem solving', 'analytical',
            'creative', 'innovative', 'collaborative', 'adaptable', 'detail-oriented',
            'self-motivated', 'organized', 'time management', 'critical thinking'
        ]
        
        for term in soft_skills_terms:
            if term in text_lower:
                keywords['soft_skills'].append({
                    'term': term,
                    'confidence': 0.8 if f' {term} ' in text_lower else 0.6
                })
        
        industry_terms = [
            'technology', 'healthcare', 'finance', 'education', 'retail', 'manufacturing',
            'consulting', 'government', 'nonprofit', 'startup', 'enterprise'
        ]
        
        for term in industry_terms:
            if term in text_lower:
                keywords['industries'].append({
                    'term': term,
                    'confidence': 0.7
                })
        
        job_titles = [
            'developer', 'engineer', 'analyst', 'manager', 'director', 'architect',
            'consultant', 'specialist', 'coordinator', 'administrator', 'designer'
        ]
        
        for title in job_titles:
            if title in text_lower:
                keywords['job_titles'].append({
                    'term': title,
                    'confidence': 0.8
                })
        
        action_words = [
            'developed', 'implemented', 'designed', 'created', 'built', 'managed',
            'led', 'optimized', 'improved', 'increased', 'reduced', 'achieved'
        ]
        
        for word in action_words:
            if word in text_lower:
                keywords['action_words'].append({
                    'term': word,
                    'confidence': 0.9
                })
        
        return keywords
    
    def _extract_achievements(self, text: str) -> List[Dict[str, Any]]:
        achievements = []
        
        achievement_matches = self.patterns['achievements'].findall(text)
        
        for match in achievement_matches:
            achievements.append({
                'text': match,
                'type': 'quantified',
                'confidence': 0.8
            })
        
        achievement_indicators = [
            r'(promoted to|advancement to|selected for|chosen for|awarded|recognized for)([^.!?]*)',
            r'(led team of|managed team of|supervised)([^.!?]*)',
            r'(exceeded expectations|surpassed goals|beat targets)([^.!?]*)',
            r'(first place|top performer|highest rated|best in)([^.!?]*)'
        ]
        
        for pattern in achievement_indicators:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                achievements.append({
                    'text': match.group(0),
                    'type': 'qualitative',
                    'confidence': 0.7
                })
        
        return achievements
    
    def _generate_document_insights(self, extraction_result: Dict[str, Any]) -> Dict[str, Any]:
        insights = {
            'document_type': self._classify_document_type(extraction_result),
            'completeness_score': 0,
            'professional_level': 'unknown',
            'key_strengths': [],
            'potential_improvements': [],
            'relevant_sections': []
        }
        
        sections = extraction_result.get('sections', {})
        contact_info = extraction_result.get('contact_info', {})
        structured_data = extraction_result.get('structured_data', {})
        keywords = extraction_result.get('keywords', {})
        
        completeness_factors = {
            'has_contact': bool(contact_info.get('emails')),
            'has_experience': 'experience' in sections,
            'has_education': 'education' in sections,
            'has_skills': 'skills' in sections,
            'has_technical_skills': bool(keywords.get('technical_skills', [])),
            'has_achievements': bool(extraction_result.get('achievements', [])),
        }
        
        insights['completeness_score'] = sum(completeness_factors.values()) / len(completeness_factors)
        
        experience_indicators = len(structured_data.get('experience', {}).get('date_ranges', []))
        tech_skills_count = len(keywords.get('technical_skills', []))
        
        if experience_indicators >= 3 and tech_skills_count >= 5:
            insights['professional_level'] = 'senior'
        elif experience_indicators >= 1 and tech_skills_count >= 3:
            insights['professional_level'] = 'mid-level'
        elif tech_skills_count >= 1:
            insights['professional_level'] = 'junior'
        else:
            insights['professional_level'] = 'entry-level'
        
        if tech_skills_count > 8:
            insights['key_strengths'].append('Strong technical background')
        if len(extraction_result.get('achievements', [])) > 2:
            insights['key_strengths'].append('Demonstrated achievements')
        if 'summary' in sections and len(sections['summary']) > 100:
            insights['key_strengths'].append('Clear professional summary')
        
        if not contact_info.get('emails'):
            insights['potential_improvements'].append('Add contact information')
        if 'skills' not in sections:
            insights['potential_improvements'].append('Add skills section')
        if not extraction_result.get('achievements'):
            insights['potential_improvements'].append('Include quantifiable achievements')
        
        if insights['document_type'].startswith('resume'):
            insights['relevant_sections'] = ['experience', 'education', 'skills', 'projects']
        elif insights['document_type'].startswith('job_description'):
            insights['relevant_sections'] = ['requirements', 'responsibilities', 'benefits']
        
        return insights
    
    def _calculate_quality_score(self, result: Dict[str, Any]) -> float:
        score = 0.0
        max_score = 100.0
        
        text_length = len(result.get('cleaned_text', ''))
        if text_length > 1000:
            score += 25
        elif text_length > 500:
            score += 20
        elif text_length > 200:
            score += 15
        elif text_length > 50:
            score += 10
        
        sections = result.get('sections', {})
        metadata_sections = sections.get('_metadata', {})
        section_count = metadata_sections.get('total_sections', len(sections))
        
        if section_count >= 5:
            score += 25
        elif section_count >= 3:
            score += 20
        elif section_count >= 2:
            score += 15
        elif section_count >= 1:
            score += 10
        
        contact = result.get('contact_info', {})
        contact_score = 0
        if contact.get('emails'):
            contact_score += 8
        if contact.get('phones'):
            contact_score += 6
        if contact.get('linkedin') or contact.get('github'):
            contact_score += 4
        if contact.get('location'):
            contact_score += 2
        score += contact_score
        
        structured = result.get('structured_data', {})
        structured_score = 0
        
        if structured.get('education', {}).get('degrees'):
            structured_score += 4
        if structured.get('experience', {}).get('date_ranges'):
            structured_score += 4
        if structured.get('technical_info', {}).get('programming_languages'):
            structured_score += 3
        if structured.get('technical_info', {}).get('frameworks'):
            structured_score += 2
        if structured.get('technical_info', {}).get('tools'):
            structured_score += 2
        score += structured_score
        
        keywords = result.get('keywords', {})
        achievements = result.get('achievements', [])
        
        tech_skills_count = len(keywords.get('technical_skills', []))
        if tech_skills_count >= 5:
            score += 4
        elif tech_skills_count >= 3:
            score += 3
        elif tech_skills_count >= 1:
            score += 2
        
        if len(achievements) >= 3:
            score += 3
        elif len(achievements) >= 1:
            score += 2
        
        soft_skills_count = len(keywords.get('soft_skills', []))
        if soft_skills_count >= 3:
            score += 3
        elif soft_skills_count >= 1:
            score += 2
        
        insights = result.get('document_insights', {})
        if insights.get('completeness_score', 0) > 0.7:
            score += 3
        elif insights.get('completeness_score', 0) > 0.5:
            score += 2
        elif insights.get('completeness_score', 0) > 0.3:
            score += 1
        
        if insights.get('professional_level') != 'unknown':
            score += 2
        
        return min(score, max_score)
    
    def get_summary(self, extraction_result: Dict[str, Any]) -> Dict[str, Any]:
        if not extraction_result.get('success', False):
            return {'error': 'Extraction failed', 'summary': None}
        
        insights = extraction_result.get('document_insights', {})
        keywords = extraction_result.get('keywords', {})
        achievements = extraction_result.get('achievements', [])
        
        summary = {
            'document_analysis': {
                'type': insights.get('document_type', 'unknown'),
                'professional_level': insights.get('professional_level', 'unknown'),
                'quality_score': extraction_result.get('quality_score', 0),
                'completeness_score': insights.get('completeness_score', 0)
            },
            'content_statistics': {
                'character_count': len(extraction_result.get('raw_text', '')),
                'word_count': len(extraction_result.get('cleaned_text', '').split()),
                'line_count': len(extraction_result.get('raw_text', '').split('\n')),
                'sections_found': len(extraction_result.get('sections', {})),
                'achievements_count': len(achievements),
                'technical_skills_count': len(keywords.get('technical_skills', []))
            },
            'key_information': {
                'has_contact_info': bool(extraction_result.get('contact_info', {}).get('emails')),
                'has_experience_section': 'experience' in extraction_result.get('sections', {}),
                'has_education_section': 'education' in extraction_result.get('sections', {}),
                'has_skills_section': 'skills' in extraction_result.get('sections', {}),
                'has_quantified_achievements': any(a.get('type') == 'quantified' for a in achievements),
                'has_professional_summary': 'summary' in extraction_result.get('sections', {})
            },
            'technical_profile': {
                'programming_languages': [skill['term'] for skill in keywords.get('technical_skills', []) 
                                        if skill.get('category') == 'programming'],
                'frameworks': [skill['term'] for skill in keywords.get('technical_skills', []) 
                             if skill.get('category') == 'frameworks'],
                'tools_and_platforms': [skill['term'] for skill in keywords.get('technical_skills', []) 
                                      if skill.get('category') in ['tools', 'cloud', 'databases']],
                'soft_skills': [skill['term'] for skill in keywords.get('soft_skills', [])]
            },
            'recommendations': {
                'strengths': insights.get('key_strengths', []),
                'improvements': insights.get('potential_improvements', []),
                'missing_sections': [section for section in insights.get('relevant_sections', []) 
                                   if section not in extraction_result.get('sections', {})]
            },
            'extraction_metadata': {
                'file_type': extraction_result.get('file_type'),
                'extraction_method': extraction_result.get('extraction_method'),
                'processing_success': extraction_result.get('success', False),
                'confidence_scores': extraction_result.get('sections', {}).get('_metadata', {}).get('section_confidence', {})
            }
        }
        
        return summary
    
    def _classify_document_type(self, extraction_result: Dict[str, Any]) -> str:
        sections = extraction_result.get('sections', {})
        contact_info = extraction_result.get('contact_info', {})
        structured_data = extraction_result.get('structured_data', {})
        
        resume_indicators = 0
        if 'experience' in sections:
            resume_indicators += 2
        if 'education' in sections:
            resume_indicators += 2
        if 'skills' in sections:
            resume_indicators += 1
        if contact_info.get('emails') or contact_info.get('phones'):
            resume_indicators += 1
        if structured_data.get('technical_info', {}).get('programming_languages'):
            resume_indicators += 1
        
        job_indicators = 0
        text_lower = extraction_result.get('cleaned_text', '').lower()
        job_keywords = ['requirements', 'qualifications', 'responsibilities', 'salary', 'benefits', 'apply', 'position', 'role']
        for keyword in job_keywords:
            if keyword in text_lower:
                job_indicators += 1
        
        if resume_indicators >= 4:
            return 'resume'
        elif job_indicators >= 3:
            return 'job_description'
        elif resume_indicators >= 2:
            return 'resume_partial'
        elif job_indicators >= 1:
            return 'job_description_partial'
        else:
            return 'unknown'